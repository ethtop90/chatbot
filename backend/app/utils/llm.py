import sys
import re
import os
import time
from datetime import datetime
from dotenv import load_dotenv
from langchain_chroma import Chroma
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.document_loaders import AsyncHtmlLoader
from langchain_community.document_transformers import Html2TextTransformer
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_community.document_transformers import LongContextReorder
import csv
import pandas as pd
import docx
from langchain.docstore.document import Document
from PyPDF2 import PdfFileReader
from app.models.learningData import LearningData  # Assuming this is the correct import path for LearningData
from app.models.user import User

load_dotenv(override=True)
OPENAI__API__KEY = os.environ.get("OEPNAI_API_KEY")

today = datetime.now()
formatted_date = today.strftime("%Y%m%d")

reordering = LongContextReorder()

llm = ChatOpenAI(model_name="gpt-3.5-turbo", openai_api_key=OPENAI__API__KEY, temperature=0.3)


def create_doc(page_content, metadata):
    doc =  Document(page_content=page_content, metadata=metadata)
    return doc

def load_file_contents(files):
    contents = []
    for file in files:
        file_path = file['path']
        file_name = file['filename']
        metadata = {"source": file_name}
        
        extension = os.path.splitext(file_path)[1].lower()
        if extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                contents.append(create_doc(f.read(), metadata))
        elif extension == '.pdf':
            contents.append(create_doc(extract_text_from_pdf(file_path), metadata))
        elif extension == '.csv':
            contents.append(create_doc(extract_text_from_csv(file_path), metadata))
        elif extension in ['.xls', '.xlsx']:
            contents.append(create_doc(extract_text_from_excel(file_path), metadata))
        elif extension == '.docx':
            contents.append(create_doc(extract_text_from_docx(file_path), metadata))
        elif extension == '.html':
            contents.append(create_doc(extract_text_from_html(file_path), metadata))
    return contents

def extract_text_from_pdf(file_path):
    text = []
    with open(file_path, 'rb') as f:
        pdf = PdfFileReader(f)
        for page_num in range(pdf.getNumPages()):
            page = pdf.getPage(page_num)
            text.append(page.extract_text())
    return ' '.join(text)

def extract_text_from_csv(file_path):
    rows = []
    with open(file_path, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(' '.join(row))
    return ' '.join(rows)

def extract_text_from_excel(file_path):
    df = pd.read_excel(file_path)
    return df.to_string()

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_text_from_html(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    html_transformer = Html2TextTransformer()
    return html_transformer.transform(html_content)

def sanitize_chatbot_id(chatbot_id):
    # Replace '@' with an underscore and remove any invalid characters
    sanitized_id = re.sub(r'[^a-zA-Z0-9_-]', '_', chatbot_id)
    # Ensure the collection name is between 3 and 63 characters
    return sanitized_id[:63]

def create_doc_from_hand_input(hand_inputs):
    contents = []
    for hand_input in hand_inputs:
        contents.append(create_doc(hand_input['content'], {"title": hand_input['title']}))
    return contents


def create_vector_db(chatbot_id):
    # Load necessary assets using LearningData model
    urls = LearningData.read_url_data(chatbot_id)
    files = LearningData.read_file_data(chatbot_id)
    hand_inputs = LearningData.read_hand_input_data(chatbot_id)

    # Initialize the loader with URLs
    url_list = [url['URL'] for url in urls]
    loader = AsyncHtmlLoader(url_list, verify_ssl=False)

    # Load HTML content from URLs
    txt = loader.load()

    # Transform loaded HTML content into plain text
    html2text = Html2TextTransformer()
    docs_transformed = html2text.transform_documents(txt)

    # Load and transform file content based on file type
    file_contents = load_file_contents(files)

    # Combine all text contents
    combined_texts = docs_transformed + file_contents + create_doc_from_hand_input(hand_inputs)

    # Split combined text into chunks
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunked_text = text_splitter.split_documents(combined_texts)

    converted_chunked_text = [''.join(str(item) for item in tup) for tup in chunked_text]

    sanitized_chatbot_id = sanitize_chatbot_id(chatbot_id)

    # Generate embeddings for chunks using Chroma and OpenAI embeddings
    vectordb = Chroma.from_texts(
        converted_chunked_text,
        embedding=OpenAIEmbeddings(openai_api_key=OPENAI__API__KEY),
        collection_name=f"{sanitized_chatbot_id}",
        persist_directory="chroma_db"
    )
    print('Size of vectorDB: {} bytes'.format(sys.getsizeof(vectordb)))
    return vectordb

def create_rag_chain(chatbot_id):
    sanitized_chatbot_id = sanitize_chatbot_id(chatbot_id)

    vectorstore = Chroma(
        persist_directory="./chroma_db",
        embedding_function=OpenAIEmbeddings(openai_api_key=OPENAI__API__KEY),
        collection_name=f"{sanitized_chatbot_id}"
    )
    
    retriever = vectorstore.as_retriever(search_kwargs={"k": 6})
    reordering.transform_documents(retriever)

    # llm = ChatOpenAI(model_name="gpt-3.5-turbo", openai_api_key=OPENAI__API__KEY, temperature=0.3)

    contextualize_q_system_prompt = f"""
    今日の日付は{today.year}年{today.month}月{today.day}日です。
    千葉県に関する行政情報や旅行情報に詳しいアシスタントです。
    ステップバイステップで考えてみましょう。ユーザーの質問を独立した形で再構築し、明確に理解できるようにしてください。
    簡潔で情報量のあるスタイルを保ち、読みやすいようにMarkdown形式で回答をフォーマットしてください。
    """
    
    contextualize_q_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", contextualize_q_system_prompt),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{question}"),
        ]
    )
    
    contextualize_q_chain = contextualize_q_prompt | llm | StrOutputParser()
    
    def contextualized_question(input: dict):
        if input.get("chat_history"):
            return contextualize_q_chain
        else:
            return input["question"]
    
    qa_system_prompt = f"""今日の日付が{today.year}年{today.month}月{today.day}日であることを覚えておいてください。"""
    qa_system_prompt = qa_system_prompt + """
    最後の質問に答えるために、以下のすべての文脈を文書で使用する。
    チャット履歴からユーザーの質問に答えることができる場合は、提供された文書を参照する必要はありません。
    提供された文書に答えがない場合は、あなたの既存の知識を活用しても大丈夫です。
    しかし、この場合には必ず回答の出所を明らかにしなければなりません。
    回答はMarkdownでフォーマットしてください。
    {context} 
    """
    
    qa_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", qa_system_prompt),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{question},Reply with japanese language"),
        ]
    )
    
    rag_chain = (
        RunnablePassthrough.assign(
            context=contextualized_question | retriever
        )
        | qa_prompt
        | llm
        | StrOutputParser()
    )
    print('Size of rag_chain: {} bytes'.format(sys.getsizeof(rag_chain)))
    return rag_chain

def create_follow_up_chain(rag_chain):
    follow_up_q_prompt = """
    ユーザーの質問に基づいて、千葉県に関連する適切なフォローアップ質問を2-4つ提案します。
    これらの質問を回答を提供せずに、ユーザーの興味に合わせてリスト形式で提示してください。重複や些細な質問を避けてください。
    10文字程度を超えない一般的なフォローアップの質問であること。
    1つの質問は10文字を超えてはならない。
    以下のフォーマットに従ってください:
    --- フォーマットテンプレート ---
    [
    "質問1",
    "質問2",
    "質問3",
    "質問4"
    ]
    """
   
    follow_up_prompt_template = ChatPromptTemplate.from_messages(
        [
            ("system", follow_up_q_prompt),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "与えられた質問 {question} に基づいて、元の質問の文脈に沿った千葉県に関する2～4個のフォローアップ質問を生成してください。質問はリスト形式で回答を含めずに提示してください。フォローアップ質問は元の質問で言及されたトピックに関連するものであり、重複や些細な質問は避けてください。日本語で返信してください。"),
        ]
    )
    
    follow_up_chain = follow_up_prompt_template | llm | StrOutputParser()
    print('Size of follow_up_chain: {} bytes'.format(sys.getsizeof(follow_up_chain)))

    return follow_up_chain

def create_keyword_chain(rag_chain):
    keyword_q_prompt = """
    あなたが学習した資料に基づいて5～10個のキーワードを生成してください。
    重複したり、些細なキーワードは避けてください。
    例："WEBサイト制作について", "代表取締役について", "キャンペーン実施中", "サービス料金を知る", "採用募集中！", "採用について", ...
    以下のフォーマットに従ってください。
    --- 形式テンプレート ---
    [
    "キーワード1",
    "キーワード2",
    "キーワード3",
    "キーワード4"
    ]
    """
   
    keyword_prompt_template = ChatPromptTemplate.from_messages(
        [
            ("system", keyword_q_prompt),
            ("human", "{user_input}"),
        ]
    )
   
    keyword_chain = keyword_prompt_template | rag_chain | StrOutputParser()
    print('Size of keyword_chain: {} bytes'.format(sys.getsizeof(keyword_chain)))

    return keyword_chain

def generate_response(rag_chain, question, chat_history=[]):
    for chunk in rag_chain.stream({'question': question, 'chat_history': chat_history}):
        yield chunk

def generate_follow_up_question(follow_up_chain, question, chat_history=[]):
    follow_up_questions = follow_up_chain.invoke({'question': question, 'chat_history': chat_history})
    return follow_up_questions

# def generate_keyword(keyword_chain, user_input="学習した資料に基づき、キーワードはリスト形式で提示してください。キーワードは、元の質問で言及されたトピックに関連するもので、重複や些細なキーワードは避けてください。 回答は日本語でお願いします。"):
#     keywords = keyword_chain.invoke({'user_input': user_input})
#     return keywords

def generate_keyword(rag_chain, user_input="""あなたが学習した資料に基づいて5～10個のキーワードを生成してください。
重複したり、些細なキーワードは避けてください。
例："WEBサイト制作について", "代表取締役について", "キャンペーン実施中", "サービス料金を知る", "採用募集中！", "採用について", ...
以下のフォーマットに従ってください。
--- 形式テンプレート ---
[
"キーワード1",
"キーワード2",
"キーワード3",
"キーワード4"
]""", chat_history=[]):
    keywords = rag_chain.invoke({'question': user_input, 'chat_history': chat_history})
    return keywords

# Example usage:
# chatbot_id = "fullsuccess.world@gmail.com"
# vectordb = create_vector_db(chatbot_id)
# rag_chain = create_rag_chain(chatbot_id)
# follow_up_chain = create_follow_up_chain(rag_chain)
# response_chunks = generate_response(rag_chain, "What is the weather like in Chiba?")
# follow_up_questions = generate_follow_up_question(follow_up_chain, "Tell me more about Chiba.")

def have_learning_data(chatbot_id):
    urls = LearningData.read_url_data(chatbot_id)
    files = LearningData.read_file_data(chatbot_id)
    hand_inputs = LearningData.read_hand_input_data(chatbot_id)
    return len(urls) + len(files) + len(hand_inputs)


chatbot_set = {}

def prepare_llm():
    users = User.get_all_users()
    for user in users:
        email = user.get('email')
        if have_learning_data(user.get('email')) == 0:
            return
        chatbot = {}
        create_vector_db(email)
        rag_chain = create_rag_chain(email)
        follow_up_chain = create_follow_up_chain(rag_chain)
        keyword_chain = create_keyword_chain(rag_chain)
        chatbot['rag_chain'] = rag_chain
        chatbot['follow_up_chain'] = follow_up_chain
        chatbot_set[email] = chatbot

def edit_llm(chatbot_id):
    if(chatbot_set.get(chatbot_id)):
        create_vector_db(chatbot_id)
        chatbot = {}
        rag_chain = create_rag_chain(chatbot_id)
        follow_up_chain = create_follow_up_chain(rag_chain)
        keyword_chain = create_keyword_chain(rag_chain)
        chatbot['rag_chain'] = rag_chain
        print("chatbot['rag_chain']", chatbot['rag_chain'])
        chatbot['follow_up_chain'] = follow_up_chain
        chatbot_set[chatbot_id] = chatbot

